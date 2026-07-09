# import docx  # यो माथि थप्नुहोस्
# from django.db import models
# from courses.models import Course
# from videos.models import Video

# class Note(models.Model):
#     title = models.CharField(max_length=255)
#     course = models.ForeignKey(Course, on_delete=models.CASCADE, related_name='notes')
#     video = models.ForeignKey(Video, on_delete=models.CASCADE, related_name='notes', null=True, blank=True) 
#     file = models.FileField(upload_to='notes/')
#     content = models.TextField(null=True, blank=True, help_text="AI ले पढ्नको लागि यहाँ नोटको टेक्स्ट पेस्ट गर्नुहोस्")
#     uploaded_at = models.DateTimeField(auto_now_add=True)

#     # यो नयाँ 'save' मेथड थप्नुहोस् जसले गर्दा फाइलबाट टेक्स्ट आफैं तानिन्छ
#     def save(self, *args, **kwargs):
#         if self.file and self.file.name.endswith('.docx') and not self.content:
#             try:
#                 doc = docx.Document(self.file)
#                 full_text = [para.text for para in doc.paragraphs]
#                 self.content = "\n".join(full_text)
#             except Exception as e:
#                 print(f"Error reading docx: {e}")
#         super().save(*args, **kwargs)

#     def __str__(self):
#         return self.title
import docx
import os
from django.db import models
from courses.models import Course
from videos.models import Video


class Note(models.Model):
    title = models.CharField(max_length=255)
    course = models.ForeignKey(Course, on_delete=models.CASCADE, related_name='notes')
    video = models.ForeignKey(Video, on_delete=models.CASCADE, related_name='notes', null=True, blank=True)
    file = models.FileField(upload_to='notes/')
    content = models.TextField(null=True, blank=True, help_text="AI ले पढ्नको लागि यहाँ नोटको टेक्स्ट पेस्ट गर्नुहोस्")
    uploaded_at = models.DateTimeField(auto_now_add=True)

    def extract_text_from_docx(self):
        """Word document बाट text तान्ने"""
        try:
            doc = docx.Document(self.file)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"⚠️ Error reading docx: {e}")
            return None

    def extract_text_from_pdf(self):
        """PDF बाट text तान्ने"""
        try:
            from pypdf import PdfReader
            self.file.seek(0)
            reader = PdfReader(self.file)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except ImportError:
            print("⚠️ pypdf not installed. Run: pip install pypdf")
            return None
        except Exception as e:
            print(f"⚠️ Error reading PDF: {e}")
            return None

    def extract_text_from_txt(self):
        """Plain text file बाट content पढ्ने"""
        try:
            self.file.seek(0)
            return self.file.read().decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"⚠️ Error reading txt: {e}")
            return None

    def save(self, *args, **kwargs):
        # ✅ FIX: content खाली छ भने मात्र, र file को extension अनुसार सही reader use गर्ने
        if self.file and not self.content:
            filename = self.file.name.lower()
            extracted = None

            if filename.endswith('.docx'):
                extracted = self.extract_text_from_docx()
            elif filename.endswith('.pdf'):
                extracted = self.extract_text_from_pdf()
            elif filename.endswith('.txt'):
                extracted = self.extract_text_from_txt()
            else:
                print(f"⚠️ Unsupported file type for auto-extraction: {filename}")
                print("   Supported: .docx, .pdf, .txt — content field मा manually paste गर्नुहोस्।")

            if extracted:
                self.content = extracted
                print(f"✅ Auto-extracted {len(extracted)} characters from {filename}")

        super().save(*args, **kwargs)

    def __str__(self):
        return self.title